from __future__ import annotations

from io import BytesIO
import json
from typing import Any

from orchid_ai.core.tool import OrchidTool, OrchidToolInput, OrchidToolOutput

from .write_file import _ensure_extension, _write_content


def _add_markdown_like_content(document: Any, content: str) -> None:
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)


def _add_sections(document: Any, sections: list[Any]) -> None:
    for section in sections:
        if isinstance(section, dict):
            title = str(section.get("title") or "Section")
            document.add_heading(title, level=2)
            duration = section.get("duration")
            if duration not in (None, ""):
                document.add_paragraph(f"Duration: {duration}")
            body = section.get("content") or section.get("content_template") or section.get("summary") or ""
            if isinstance(body, list):
                for item in body:
                    document.add_paragraph(str(item), style="List Bullet")
            elif isinstance(body, dict):
                document.add_paragraph(json.dumps(body, indent=2, sort_keys=True))
            elif body:
                _add_markdown_like_content(document, str(body))
        else:
            document.add_paragraph(str(section))


class GenerateDOCXTool(OrchidTool):
    name = "generate_docx"
    description = "Generate a DOCX file from structured or Markdown-like content"
    parameters_schema = {
        "type": "object",
        "properties": {
            "content": {
                "type": "string",
                "description": "Document content in text or Markdown-like form",
            },
            "filename": {
                "type": "string",
                "description": "Output filename without extension",
            },
            "title": {
                "type": "string",
                "description": "Optional document title",
                "default": "",
            },
            "sections": {
                "type": "array",
                "description": "Optional structured sections",
                "default": [],
            },
        },
        "required": ["content", "filename"],
    }
    parallel_safe = False

    async def invoke(self, tool_input: OrchidToolInput) -> OrchidToolOutput:
        from docx import Document

        parameters = tool_input.parameters
        document = Document()

        title = str(parameters.get("title", "") or "")
        if title:
            document.core_properties.title = title
            document.add_heading(title, level=0)

        sections = parameters.get("sections") or []
        if isinstance(sections, list) and sections:
            _add_sections(document, sections)
        else:
            _add_markdown_like_content(document, str(parameters["content"]))

        stream = BytesIO()
        document.save(stream)
        filepath = str(_ensure_extension(str(parameters["filename"]), ".docx"))
        path, size, download_url = _write_content(
            stream.getvalue(),
            filepath=filepath,
            mode="binary",
            context=tool_input.context,
            content_sources=tool_input.content_sources,
            auth_context=tool_input.auth_context,
        )
        return OrchidToolOutput(
            result={"path": str(path), "size_bytes": size, "format": "docx", "download_url": download_url},
            metadata={
                "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            },
        )
